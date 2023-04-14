#import libraries
import openpyxl
import docx

# total fees collected
def calculateTotalFees(sheet, dueStudents):
     totalStudents = sheet.max_row+1
     paidStudents = totalStudents - len(dueStudents)
     #1500 dollar is fees for one student.
     totalFees = paidStudents * 1500

     return totalFees
     

try:
     
     #load the student and expense excel files
     studentList = openpyxl.load_workbook("Pythonproject.xlsx")
     sheet = studentList.active
     columnList1 = ["A","B","C"]
     
     expenseList = openpyxl.load_workbook("Expense.xlsx")
     expenseSheet = expenseList.active
     columnList2 = ["A","B"]
     
     
     #calculate total expense
     totalExpenses = 0
     for row in range(2,expenseSheet.max_row + 1):
          expenseAmountNum = expenseSheet[str(columnList2[1]) + str(row)]
          expenseAmount = expenseAmountNum.value
          if expenseAmount:
            totalExpenses += int(expenseAmount)
            

     #student who have not paid the fees
     dueStudents = []
     row = 2
     while row <= sheet.max_row:
          studentCellNumber = sheet[str(columnList1[0]) + str(row)]
          studentName = str(studentCellNumber.value)
          feeStatusCellNum = sheet[str(columnList1[1]) + str(row)]
          feesStatus = str(feeStatusCellNum.value)
          phoneNumberCell = sheet[str(columnList1[2]) + str(row)]
          phoneNumber = str(phoneNumberCell.value)                        
          if feesStatus == "Not Paid":
               dueStudents.append((studentName, phoneNumber))
          row = row + 1
          

     #total fees collected and profit or loss calculated
     totalFees =  calculateTotalFees(sheet, dueStudents)

     profitLossAmount = totalFees - totalExpenses
     if profitLossAmount >= 0:
          profitLoss = "Profit of $" + str(profitLossAmount)
     else:
          profitLoss = "Loss of $" + str(profitLossAmount)
          

     #creating a new word document and add headings and paragraphs                         
     document = docx.Document()
     document.add_heading("List of Students with Due Fees")
     for i in range(len(dueStudents)):
         student = dueStudents[i]
         document.add_paragraph(str(i+1) + '. ' + student[0] + ' (' + student[1] + ')')
         
     document.add_heading("Financial Report")
     document.add_paragraph("Total Expenses:$ " + str(totalExpenses))
     document.add_paragraph("Total Fees Collected:$ " + str(totalFees))
     document.add_paragraph("Profit/Loss: " + str(profitLoss))
     
     #saving the word document
     document.save("Studentdues.docx")

     
except Exception as e:
     #catch the errors that occur during execution.
     print("An error occurred:", e)
